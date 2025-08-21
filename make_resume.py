from docx import Document
from docx.shared import Pt

document = Document()

# Name & Contact Info
name = document.add_paragraph("Parth Jatin Shah")
name.runs[0].bold = True
name.runs.font.size = Pt(16)

document.add_paragraph("Panvel, Navi Mumbai | +91 9834343014 | parthjhk@gmail.com")
document.add_paragraph("LinkedIn: https://www.linkedin.com/in/parth-shah-93a014316/")
document.add_paragraph("GitHub: https://github.com/ParthShah007")

# Section Helper
def add_section(title, content_list):
    document.add_paragraph("")
    heading = document.add_paragraph(title)
    heading.runs.bold = True
    heading.runs.underline = True
    for item in content_list:
        document.add_paragraph("• " + item, style='List Bullet')

# Summary
document.add_paragraph("\nSummary", 'Heading 2')
document.add_paragraph(
    "Passionate about Machine Learning and AI, with strong foundations in Python, Data Analysis, and Problem-solving. "
    "Enthusiastic about exploring new technologies and applying ML/AI models to solve real-world challenges."
)

# Education
document.add_paragraph("\nEducation", 'Heading 2')
document.add_paragraph(
    "DJ Sanghvi College of Engineering – B.Tech in Computer Science & Engineering (Data Science)\n"
    "2024 – 2028 | 1st Year Avg SGPA: 9.57"
)

# Skills
add_section("Technical Skills", [
    "Programming: Python, Java, C",
    "Data & AI: Pandas, NumPy, Scikit-Learn, Matplotlib, Seaborn",
    "Tools: Jupyter, VS Code, GitHub, Anaconda",
    "Databases: PostgreSQL"
])

# Projects
add_section("Projects", [
    "File Manager (Python): Created using file handling and pathlib.",
    "Car & Insurance Data Modeling: Built ML models on Ford dataset and Kaggle insurance dataset.",
    "Bike Rental Forecasting: Applied regression techniques to forecast daily rentals."
])

# Achievements
add_section("Achievements", [
    "Ranked 19th in SPIT CodeBuster Event",
    "Kaggle Pandas Certification",
    "Solved 80+ LeetCode problems",
    "Hackerrank Python – 5⭐"
])

# Hobbies
add_section("Hobbies", [
    "Pro-level Typeracer", "Chess", "Football", "Strategy Games"
])

document.save("Resume_Parth_Shah.docx")
print("✅ Resume saved as Resume_Parth_Shah.docx")
